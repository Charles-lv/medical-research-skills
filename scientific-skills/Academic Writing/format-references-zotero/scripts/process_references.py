#!/usr/bin/env python3
"""
process_references.py - Zotero version
Convert [PMID: xxxx] markers to Word doc with native Zotero field codes.

Usage: python process_references.py <markdown_file> [--style <name>]

Style defaults to 'nature'. Common values: nature, vancouver, apa, nejm, lancet

Pipeline:
  1. Extract & deduplicate PMIDs
  2. Fetch PubMed metadata, convert to CSL JSON
  3. Replace [PMID: xxxx] with ZRCITE markers, compile .docx via Pandoc
  4. Replace ZRCITE markers with ADDIN ZOTERO_ITEM field codes (python-docx)
  5. Append ADDIN ZOTERO_BIBL and ZOTERO_PREF fields
  6. Write <stem>.ris for Zotero library import
  7. Open both files
"""

import sys
import re
import os
import copy
import json
import time
import html
import zipfile
import shutil
import platform
import subprocess
import urllib.request
import urllib.error
import ssl
import xml.etree.ElementTree as ET
from pathlib import Path

ssl._create_default_https_context = ssl._create_unverified_context

try:
    from docx import Document
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print('Missing python-docx, please run: pip install python-docx', file=sys.stderr)
    sys.exit(1)


def resolve_style(name: str) -> str:
    if name.startswith('http'):
        return name
    return f'http://www.zotero.org/styles/{name}'


# ── PubMed fetch ─────────────────────────────────────────────────────────────

def extract_pmids(text: str) -> list:
    found = re.findall(r'\[PMID:\s*(\d+)\]', text, re.IGNORECASE)
    seen = []
    for p in found:
        if p not in seen:
            seen.append(p)
    return seen


def _text(el):
    if el is None:
        return ''
    parts = [el.text or '']
    for child in el:
        parts.append(_text(child))
        parts.append(child.tail or '')
    return ''.join(parts).strip()


def _transliterate_greek(text: str) -> str:
    """Replace Greek Unicode letters with Latin equivalents.

    Shared normalization used by both EndNote and Zotero scripts so that
    the .ris file and shared PubMed cache are always consistent.
    """
    _GREEK_MAP = {
        '\u03b1': 'alpha', '\u03b2': 'beta', '\u03b3': 'gamma',
        '\u03b4': 'delta', '\u03b5': 'epsilon', '\u03b6': 'zeta',
        '\u03b7': 'eta', '\u03b8': 'theta', '\u03b9': 'iota',
        '\u03ba': 'kappa', '\u03bb': 'lambda', '\u03bc': 'mu',
        '\u03bd': 'nu', '\u03be': 'xi', '\u03bf': 'omicron',
        '\u03c0': 'pi', '\u03c1': 'rho', '\u03c3': 'sigma',
        '\u03c2': 'sigma',
        '\u03c4': 'tau', '\u03c5': 'upsilon', '\u03c6': 'phi',
        '\u03c7': 'chi', '\u03c8': 'psi', '\u03c9': 'omega',
        '\u0391': 'Alpha', '\u0392': 'Beta', '\u0393': 'Gamma',
        '\u0394': 'Delta', '\u0395': 'Epsilon', '\u0396': 'Zeta',
        '\u0397': 'Eta', '\u0398': 'Theta', '\u0399': 'Iota',
        '\u039a': 'Kappa', '\u039b': 'Lambda', '\u039c': 'Mu',
        '\u039d': 'Nu', '\u039e': 'Xi', '\u039f': 'Omicron',
        '\u03a0': 'Pi', '\u03a1': 'Rho', '\u03a3': 'Sigma',
        '\u03a4': 'Tau', '\u03a5': 'Upsilon', '\u03a6': 'Phi',
        '\u03a7': 'Chi', '\u03a8': 'Psi', '\u03a9': 'Omega',
    }
    result = []
    for ch in text:
        result.append(_GREEK_MAP.get(ch, ch))
    return ''.join(result)


def _fetch_xml(pmids: list) -> str:
    ids = ','.join(pmids)
    url = (
        'https://eutils.ncbi.nlm.nih.gov/entrez/eutils/efetch.fcgi'
        f'?db=pubmed&id={ids}&retmode=xml'
    )
    req = urllib.request.Request(url, headers={'User-Agent': 'format-references-zotero/1.0'})
    with urllib.request.urlopen(req, timeout=30) as resp:
        return resp.read().decode('utf-8')


def _parse_articles(xml_data: str) -> dict:
    root = ET.fromstring(xml_data)
    result = {}
    for article in root.findall('.//PubmedArticle'):
        pmid_el = article.find('.//PMID')
        if pmid_el is None:
            continue
        pmid = pmid_el.text.strip()

        author_els = article.findall('.//Author')
        first_author = 'Unknown'
        all_authors = []
        for auth in author_els:
            ln = auth.findtext('LastName', '').strip()
            fn = auth.findtext('ForeName', '').strip()
            initials = auth.findtext('Initials', '').strip()
            if ln:
                display = f'{ln}, {fn}' if fn else (f'{ln} {initials}' if initials else ln)
                all_authors.append(display)
                if first_author == 'Unknown':
                    first_author = ln

        year = 'n.d.'
        pub_date = article.find('.//PubDate')
        if pub_date is not None:
            year_el = pub_date.find('Year')
            if year_el is not None and year_el.text:
                year = year_el.text.strip()
            else:
                med = pub_date.findtext('MedlineDate', '')
                if med:
                    year = med[:4]
        if year == 'n.d.':
            ad = article.find('.//ArticleDate')
            if ad is not None:
                y = ad.findtext('Year', '')
                if y:
                    year = y.strip()

        title_el = article.find('.//ArticleTitle')
        full_title = _text(title_el) if title_el is not None else ''
        full_title_clean = _transliterate_greek(re.sub(r'[{}\[\]]', '', full_title).rstrip('.'))

        journal = article.findtext('.//Journal/Title', '').strip()
        volume  = article.findtext('.//Volume', '').strip()
        issue   = article.findtext('.//Issue', '').strip()
        pages   = article.findtext('.//MedlinePgn', '').strip()

        doi = ''
        for loc in article.findall('.//ELocationID'):
            if loc.get('EIdType') == 'doi':
                doi = (loc.text or '').strip()
                break

        abstract_parts = []
        for ab in article.findall('.//AbstractText'):
            label = ab.get('Label', '')
            txt = _text(ab)
            abstract_parts.append(f'{label}: {txt}' if label else txt)
        abstract = ' '.join(abstract_parts)[:500]

        result[pmid] = {
            'first_author': first_author,
            'year': year,
            'full_title': full_title_clean,
            'all_authors': all_authors,
            'journal': journal,
            'volume': volume,
            'issue': issue,
            'pages': pages,
            'doi': doi,
            'abstract': abstract,
        }
    return result


def _load_cache(cache_path: Path) -> dict | None:
    if cache_path.exists():
        try:
            return json.loads(cache_path.read_text(encoding='utf-8'))
        except Exception:
            return None
    return None


def _save_cache(cache_path: Path, metadata: dict):
    cache_path.write_text(json.dumps(metadata, ensure_ascii=False, indent=2), encoding='utf-8')


def fetch_pubmed_metadata(pmids: list, cache_path: Path = None) -> dict:
    metadata = {}

    # Try cache first
    if cache_path:
        cached = _load_cache(cache_path)
        if cached:
            hit = 0
            for pmid in pmids:
                if pmid in cached and not cached[pmid].get('full_title', '').startswith('[fetch failed'):
                    metadata[pmid] = cached[pmid]
                    hit += 1
            pmids = [p for p in pmids if p not in metadata]
            if hit:
                print(f'  Cache hit: {hit} articles, {len(pmids)} remaining to fetch')

    if not pmids:
        return metadata

    batch_size = 20
    for i in range(0, len(pmids), batch_size):
        batch = pmids[i:i + batch_size]
        try:
            xml_data = _fetch_xml(batch)
            metadata.update(_parse_articles(xml_data))
        except Exception as exc:
            print(f'  ⚠  PubMed fetch failed for {batch}: {exc}', file=sys.stderr)
        for pmid in batch:
            if pmid not in metadata:
                metadata[pmid] = {
                    'first_author': f'PMID{pmid}',
                    'year': 'n.d.',
                    'full_title': f'[fetch failed — PMID:{pmid}]',
                    'all_authors': [],
                    'journal': '', 'volume': '', 'issue': '',
                    'pages': '', 'doi': '', 'abstract': '',
                }
        if i + batch_size < len(pmids):
            time.sleep(0.34)

    if cache_path:
        _save_cache(cache_path, metadata)
    return metadata


# ── CSL JSON conversion ─────────────────────────────────────────────────────

def _parse_author(author_str: str) -> dict:
    """Convert 'Smith, John' or 'Smith JK' to CSL author dict."""
    if ', ' in author_str:
        parts = author_str.split(', ', 1)
        return {'family': parts[0], 'given': parts[1]}
    elif ' ' in author_str:
        parts = author_str.rsplit(' ', 1)
        return {'family': parts[0], 'given': parts[1]}
    return {'family': author_str}


def to_csl_item(pmid: str, meta: dict) -> dict:
    """Convert PubMed metadata dict to a CSL JSON item."""
    authors = [_parse_author(a) for a in meta.get('all_authors', []) if a]

    year_str = meta.get('year', '')
    try:
        issued = {'date-parts': [[int(year_str)]]}
    except (ValueError, TypeError):
        issued = {}

    item = {
        'id': pmid,
        'type': 'article-journal',
        'title': meta.get('full_title', ''),
        'author': authors,
        'issued': issued,
        'container-title': meta.get('journal', ''),
        'volume': meta.get('volume', ''),
        'issue': meta.get('issue', ''),
        'page': meta.get('pages', ''),
        'DOI': meta.get('doi', ''),
        'PMID': pmid,
    }
    return {k: v for k, v in item.items() if v not in ('', [], {}, None)}


# ── Zotero field code builders ──────────────────────────────────────────────

def build_zotero_citation_json(pmid: str, meta: dict, citation_id: int) -> dict:
    """Build the JSON blob that goes inside ADDIN ZOTERO_ITEM CSL_CITATION."""
    csl_item = to_csl_item(pmid, meta)
    first = meta.get('first_author', 'Author')
    year  = meta.get('year', 'n.d.')
    return {
        'citationID': f'zr{citation_id}',
        'properties': {
            'formattedCitation': f'({first}, {year})',
            'plainCitation': f'({first}, {year})',
            'noteIndex': 0,
        },
        'citationItems': [
            {
                'id': pmid,
                'uris': [f'http://www.ncbi.nlm.nih.gov/pubmed/{pmid}'],
                'itemData': csl_item,
            }
        ],
        'schema': 'https://raw.githubusercontent.com/citation-style-language/schema/master/csl-citation.json',
    }


def _make_field_runs(instr: str, display: str) -> list:
    """Return the 5 w:r elements that form a Word field code."""
    def r_with_fldChar(fld_type):
        r = OxmlElement('w:r')
        fc = OxmlElement('w:fldChar')
        fc.set(qn('w:fldCharType'), fld_type)
        r.append(fc)
        return r

    r_begin = r_with_fldChar('begin')

    r_instr = OxmlElement('w:r')
    it = OxmlElement('w:instrText')
    it.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    it.text = instr
    r_instr.append(it)

    r_sep = r_with_fldChar('separate')

    r_disp = OxmlElement('w:r')
    t = OxmlElement('w:t')
    t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    t.text = display
    r_disp.append(t)

    r_end = r_with_fldChar('end')

    return [r_begin, r_instr, r_sep, r_disp, r_end]


def make_citation_field_runs(citation_json: dict) -> list:
    instr = (' ADDIN ZOTERO_ITEM CSL_CITATION '
             + json.dumps(citation_json, ensure_ascii=False, separators=(',', ':'))
             + ' ')
    display = citation_json['properties']['formattedCitation']
    return _make_field_runs(instr, display)


def make_bibliography_field_runs() -> list:
    bibl = '{"uncited":[],"omittedItems":[],"custom":[]}'
    instr = f' ADDIN ZOTERO_BIBL {bibl} CSL_BIBLIOGRAPHY '
    return _make_field_runs(instr, '(Bibliography will appear here after Refresh)')


# ── Marker replacement in docx ──────────────────────────────────────────────

MARKER_PREFIX = 'ZRCITE'

def _replace_in_paragraph(para, metadata: dict, counter: list):
    """
    Find ZRCITE{pmid} tokens in a paragraph and replace with field code runs.
    counter is a single-element list used as a mutable int.
    """
    full_text = ''.join(r.text for r in para.runs)
    if MARKER_PREFIX not in full_text:
        return

    # Save rPr from first run to restore character formatting on plain text
    first_rpr = None
    if para.runs:
        rpr = para.runs[0]._r.find(qn('w:rPr'))
        if rpr is not None:
            first_rpr = copy.deepcopy(rpr)

    # Remove all existing runs
    p = para._p
    for run in list(para.runs):
        p.remove(run._r)

    pattern = re.compile(rf'({MARKER_PREFIX}\d+)')
    for part in pattern.split(full_text):
        m = re.match(rf'{MARKER_PREFIX}(\d+)', part)
        if m:
            pmid = m.group(1)
            if pmid in metadata:
                counter[0] += 1
                cj = build_zotero_citation_json(pmid, metadata[pmid], counter[0])
                for r_el in make_citation_field_runs(cj):
                    p.append(r_el)
        elif part:
            r_el = OxmlElement('w:r')
            if first_rpr is not None:
                r_el.append(copy.deepcopy(first_rpr))
            t = OxmlElement('w:t')
            t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
            t.text = part
            r_el.append(t)
            p.append(r_el)


def replace_all_markers(doc: Document, metadata: dict):
    """Walk every paragraph in the document and replace ZRCITE markers."""
    counter = [0]
    for para in doc.paragraphs:
        _replace_in_paragraph(para, metadata, counter)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    _replace_in_paragraph(para, metadata, counter)
    return counter[0]


def append_bibliography(doc: Document):
    doc.add_heading('References', level=1)
    para = doc.add_paragraph()
    p = para._p
    for r_el in make_bibliography_field_runs():
        p.append(r_el)


# ── RIS generation ──────────────────────────────────────────────────────────

def generate_ris(metadata: dict, output_path: Path):
    if output_path.exists():
        print(f'  RIS already exists, skipping (can be shared): {output_path}')
        return
    with output_path.open('w', encoding='utf-8') as f:
        for pmid, meta in metadata.items():
            f.write('TY  - JOUR\n')
            for author in meta.get('all_authors', []):
                f.write(f'AU  - {author}\n')
            f.write(f'PY  - {meta["year"]}\n')
            f.write(f'TI  - {meta.get("full_title", "")}\n')
            if meta.get('journal'):
                f.write(f'JO  - {meta["journal"]}\n')
            if meta.get('volume'):
                f.write(f'VL  - {meta["volume"]}\n')
            if meta.get('issue'):
                f.write(f'IS  - {meta["issue"]}\n')
            if meta.get('pages'):
                f.write(f'SP  - {meta["pages"]}\n')
            if meta.get('doi'):
                f.write(f'DO  - {meta["doi"]}\n')
            if meta.get('abstract'):
                f.write(f'AB  - {meta["abstract"]}\n')
            f.write(f'AN  - {pmid}\n')
            f.write('ER  - \n\n')


# ── Pandoc compile ──────────────────────────────────────────────────────────

def compile_docx(md_path: Path, docx_path: Path) -> bool:
    try:
        result = subprocess.run(
            ['pandoc', str(md_path), '-o', str(docx_path), '--from', 'markdown'],
            capture_output=True, text=True, timeout=120,
        )
        if result.returncode != 0:
            print(f'  Pandoc stderr: {result.stderr}', file=sys.stderr)
            return False
        return True
    except FileNotFoundError:
        print('\n  Pandoc not found. Please install: https://pandoc.org/installing.html\n',
              file=sys.stderr)
        return False


# ── File open ───────────────────────────────────────────────────────────────

def open_file(path: Path):
    system = platform.system()
    try:
        if system == 'Windows':
            os.startfile(str(path))
        elif system == 'Darwin':
            subprocess.run(['open', str(path)], check=True)
        else:
            subprocess.run(['xdg-open', str(path)], check=True)
    except Exception as exc:
        print(f'  Cannot auto-open {path.name}: {exc}', file=sys.stderr)
        print(f'     Please open manually: {path}', file=sys.stderr)


# ── ZOTERO_PREF injection ──────────────────────────────────────────────────

def _inject_pref(docx_path: Path, style_url: str):
    """Write ZOTERO_PREF into docx XML directly (survives Zotero Refresh)."""
    pref = {'dataVersion': 4, 'style': style_url,
            'hasBibliography': 1, 'bibliographyStyleHasBeenSet': 1}
    encoded = json.dumps(pref, ensure_ascii=False, separators=(',', ':')).replace('"', '&quot;')
    field = f'<w:p><w:fldSimple w:instr=" ADDIN ZOTERO_PREF {encoded} "/></w:p>'

    tmp = docx_path.with_suffix('.pref_tmp.docx')
    with zipfile.ZipFile(docx_path, 'r') as zin, \
         zipfile.ZipFile(tmp, 'w', zipfile.ZIP_DEFLATED) as zout:
        for item in zin.infolist():
            data = zin.read(item.filename)
            if item.filename == 'word/document.xml':
                xml = data.decode('utf-8')
                if 'ZOTERO_PREF' in xml:
                    xml = re.sub(
                        r'(ADDIN ZOTERO_PREF\s+)(\{[^<\]]*?\})',
                        lambda m: m.group(1) + json.dumps(pref, ensure_ascii=False, separators=(',', ':')).replace('"', '&quot;'),
                        xml
                    )
                else:
                    xml = xml.replace('<w:body>', '<w:body>' + field, 1)
                data = xml.encode('utf-8')
            zout.writestr(item, data)
    shutil.move(str(tmp), str(docx_path))


# ── Main ────────────────────────────────────────────────────────────────────

def main():
    args = sys.argv[1:]
    if not args:
        print('Usage: python process_references.py <markdown_file> [--style <name>]')
        sys.exit(1)

    style_name = 'nature'
    md_args = []
    force_refresh = False
    i = 0
    while i < len(args):
        if args[i] == '--style' and i + 1 < len(args):
            style_name = args[i + 1]
            i += 2
        elif args[i] == '--force-refresh':
            force_refresh = True
            i += 1
        else:
            md_args.append(args[i])
            i += 1

    if not md_args:
        print('Usage: python process_references.py <markdown_file> [--style <name>]')
        sys.exit(1)

    style_url = resolve_style(style_name)

    # ── Check / auto-install style ──
    sys.path.insert(0, str(Path(__file__).parent))
    from search_styles import find_zotero_styles_dir, install_style
    styles_dir = find_zotero_styles_dir()
    if styles_dir and not (styles_dir / f'{style_name}.csl').exists():
        print(f'Style [{style_name}] not installed, downloading...')
        if not install_style(style_name):
            print(f'Cannot install style {style_name}, please install manually and retry:')
            print(f'   python search_styles.py --install {style_name}')
            sys.exit(1)
    elif not styles_dir:
        print(f'Zotero styles directory not found, skipping style check')

    md_file = Path(md_args[0]).resolve()
    if not md_file.exists():
        print(f'File does not exist: {md_file}')
        sys.exit(1)

    out_dir = md_file.parent
    stem = md_file.stem

    # ── Step 1: read & extract PMIDs ──
    print(f'\nReading file: {md_file}')
    text = md_file.read_text(encoding='utf-8')
    pmids = extract_pmids(text)
    if not pmids:
        print('No [PMID: xxxx] markers found, nothing to process.')
        sys.exit(0)
    print(f'Found {len(pmids)} unique PMIDs')

    # Cache path: <stem>_pubmed_cache.json alongside the .md file
    cache_path = out_dir / f'{stem}_pubmed_cache.json'
    if force_refresh and cache_path.exists():
        cache_path.unlink()
        print('  Cache cleared, will re-fetch')

    # ── Step 2: fetch PubMed metadata ──
    print('Fetching metadata from PubMed...')
    metadata = fetch_pubmed_metadata(pmids, cache_path=cache_path)
    fetched = sum(1 for m in metadata.values()
                  if not m['full_title'].startswith('[fetch failed'))

    # Integrity check
    failed_pmids = [p for p in pmids
                    if metadata.get(p, {}).get('full_title', '').startswith('[fetch failed')]
    if failed_pmids:
        print(f'  {len(failed_pmids)} PMIDs failed to fetch: {", ".join(failed_pmids)}')
        print(f'     Check network and retry: python process_references.py "<file>" --force-refresh')
    else:
        print(f'Fetch complete ({fetched}/{len(pmids)} successful)')
    print(f'Fetch complete ({fetched}/{len(pmids)} successful)')

    # ── Step 3: replace markers in md, compile via Pandoc ──
    def make_marker(pmid):
        return f'{MARKER_PREFIX}{pmid}'

    marked_text = re.sub(
        r'\[PMID:\s*(\d+)\]',
        lambda m: make_marker(m.group(1)),
        text,
        flags=re.IGNORECASE,
    )
    temp_md = out_dir / f'{stem}_zr_temp.md'
    temp_md.write_text(marked_text, encoding='utf-8')

    temp_docx = out_dir / f'{stem}_zr_temp.docx'
    print('Compiling intermediate document with Pandoc...')
    if not compile_docx(temp_md, temp_docx):
        temp_md.unlink(missing_ok=True)
        sys.exit(1)
    temp_md.unlink(missing_ok=True)

    # ── Step 4: inject Zotero field codes ──
    print('Writing Zotero field codes...')
    doc = Document(str(temp_docx))
    total_refs = replace_all_markers(doc, metadata)
    append_bibliography(doc)

    docx_path = out_dir / f'{stem}-zotero.docx'
    doc.save(str(docx_path))
    temp_docx.unlink(missing_ok=True)
    _inject_pref(docx_path, style_url)
    print(f'Word document: {docx_path}')

    # ── Step 5: write RIS ──
    ris_path = out_dir / f'{stem}.ris'
    generate_ris(metadata, ris_path)
    print(f'RIS file: {ris_path}')

    # ── Step 6: open files ──
    print('\nOpening files...')
    open_file(ris_path)
    time.sleep(2)
    open_file(docx_path)

    print('\n' + '─' * 50)
    print('Done!')
    print()
    print(f'  {total_refs} citations in document')
    print(f'  Processed {len(pmids)} references')
    print(f'  Word document: {docx_path}')
    print(f'  RIS file:  {ris_path}')
    print('─' * 50)


if __name__ == '__main__':
    main()

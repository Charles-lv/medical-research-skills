#!/usr/bin/env python3
"""
Convert reference numbering style in DOCX.

Supports three scenarios:
  1. Plain text numbering: [N] <-> N. (bibliography list + inline citations changed simultaneously)
  2. Unicode superscript: [N] -> superscript characters (e.g. superscript digits), combination rules configurable
  3. Font superscript: [N] -> font format superscript (run.font.superscript=True)

Parameters (must be extracted from the actual reference article's writing style, do not assume):
  --ref-list-style    Bibliography list numbering style: bracket|period|keep
  --style             Inline citation style: keep|unicode|font
  --range-joiner      Consecutive citation range joiner (if reference article writes superscript 1-3, use --range-joiner -)
  --group-separator   Non-consecutive citation separator (if reference article writes superscript 1,3, use --group-separator ,)

Usage examples:
  # [N] -> N. (plain text)
  python convert_ref_format.py --ref-list-style period --style keep input.docx output.docx

  # [N] -> Unicode superscript, merge consecutive, comma for non-consecutive
  python convert_ref_format.py --ref-list-style keep --style unicode --range-joiner - --group-separator , input.docx output.docx

  # [N] -> Font superscript, no merging
  python convert_ref_format.py --ref-list-style keep --style font input.docx output.docx
"""
import argparse
import re
import sys
from copy import deepcopy

from docx import Document
from docx.oxml.ns import qn

# Unicode superscript mapping
SUP = {
    "0": "\u2070", "1": "\u00B9", "2": "\u00B2", "3": "\u00B3",
    "4": "\u2074", "5": "\u2075", "6": "\u2076", "7": "\u2077",
    "8": "\u2078", "9": "\u2079",
}


# ==================== Helpers ====================

def to_sup(n):
    return "".join(SUP[c] for c in str(n))


def _parse_citation_groups(text):
    """Group [N][N+1]... sequences in the body text into number lists"""
    pat = r"\[(\d+)\]"
    matches = list(re.finditer(pat, text))
    if not matches:
        return text, []
    return text, matches


# ==================== Style: bracket ↔ period ====================

def _convert_ref_list_bracket_to_period(doc):
    """Bibliography: [N] -> N."""
    count, ref_start = 0, False
    for para in doc.paragraphs:
        t = para.text.strip()
        # "References" is the Chinese heading for "References" - kept as-is because
        # this script processes Chinese-language DOCX files where the references
        # section uses this exact Chinese heading for pattern matching.
        if "References" in t and para.style.name.startswith("Heading"):
            ref_start = True
            continue
        if ref_start and t and para.runs:
            m = re.match(r"^\[(\d+)\](.*)", para.runs[0].text)
            if m:
                para.runs[0].text = f"{m.group(1)}. {m.group(2)}"
                count += 1
    return count


def _convert_ref_list_period_to_bracket(doc):
    """Bibliography: N. -> [N]"""
    count, ref_start = 0, False
    for para in doc.paragraphs:
        t = para.text.strip()
        # "References" = Chinese for "References"; used to detect the references
        # heading in Chinese-language documents processed by this script.
        if "References" in t and para.style.name.startswith("Heading"):
            ref_start = True
            continue
        if ref_start and t and para.runs:
            m = re.match(r"^(\d+)\.\s*(.*)", para.runs[0].text)
            if m:
                para.runs[0].text = f"[{m.group(1)}] {m.group(2)}"
                count += 1
    return count


# ==================== Style: superscript (unicode) ====================

def _convert_inline_to_unicode_sup(text, range_joiner, group_separator):
    """Convert inline [N] groups to Unicode superscript"""
    text, matches = _parse_citation_groups(text)
    if not matches:
        return text

    result, last_end, i = [], 0, 0
    while i < len(matches):
        result.append(text[last_end : matches[i].start()])
        nums = [int(matches[i].group(1))]
        j = i + 1
        while j < len(matches) and matches[j].start() == matches[j - 1].end():
            nums.append(int(matches[j].group(1)))
            j += 1

        if len(nums) == 1:
            result.append(to_sup(nums[0]))
        elif range_joiner is not None and all(
            nums[k + 1] - nums[k] == 1 for k in range(len(nums) - 1)
        ):
            result.append(to_sup(nums[0]) + range_joiner + to_sup(nums[-1]))
        elif group_separator is not None:
            result.append(group_separator.join(to_sup(n) for n in nums))
        else:
            result.append("".join(to_sup(n) for n in nums))

        i, last_end = j, matches[j - 1].end()
    result.append(text[last_end:])
    return "".join(result)


def _apply_unicode_superscript(doc, range_joiner, group_separator):
    """Traverse the document and replace citations with Unicode superscript"""
    ref_start = False
    count = 0
    for para in doc.paragraphs:
        t = para.text.strip()
        # "References" = Chinese for "References"; marks the start of the
        # bibliography section so inline citation conversion stops here.
        if "References" in t and para.style.name.startswith("Heading"):
            ref_start = True
        if ref_start:
            continue
        new_text = _convert_inline_to_unicode_sup(para.text, range_joiner, group_separator)
        if new_text != para.text:
            # Clear all XML content of the paragraph, preserve paragraph properties
            # Note: Cannot just use run.text = "" to clear, because DOCX XML may contain
            # hidden w:t nodes not in para.runs that run.text cannot clear
            pPr = para._element.find(qn("w:pPr"))
            pPr_copy = deepcopy(pPr) if pPr is not None else None
            para._element.clear()
            if pPr_copy is not None:
                para._element.append(pPr_copy)
            # Add new run
            from docx.oxml import OxmlElement
            new_r = OxmlElement("w:r")
            new_t = OxmlElement("w:t")
            new_t.text = new_text
            new_t.set(qn("xml:space"), "preserve")
            new_r.append(new_t)
            para._element.append(new_r)
            count += 1
    return count


# ==================== Style: superscript (font) ====================

def _split_run_for_superscript(run):
    """Split a run containing [N] into multiple runs, setting citation parts as superscript.
    Returns whether a split was made."""
    text = run.text
    if "[" not in text or "]" not in text:
        return False
    parts = re.split(r"(\[\d+\])", text)
    if len(parts) <= 1:
        return False

    parent = run._element.getparent()
    for part in parts:
        new_run = deepcopy(run)
        new_run.text = part
        rPr = new_run.find(qn("w:rPr"))
        if rPr is None:
            rPr = deepcopy(new_run.makeelement(qn("w:rPr"), {}))
            new_run.insert(0, rPr)
        if re.match(r"^\[\d+\]$", part):
            vertAlign = rPr.find(qn("w:vertAlign"))
            if vertAlign is None:
                vertAlign = deepcopy(rPr.makeelement(qn("w:vertAlign"), {}))
                rPr.append(vertAlign)
            vertAlign.set(qn("w:val"), "superscript")
        else:
            va = rPr.find(qn("w:vertAlign"))
            if va is not None:
                rPr.remove(va)
        run._element.addprevious(new_run)
    parent.remove(run._element)
    return True


def _apply_font_superscript(doc):
    """Traverse the document and split [N] citation runs into font superscript"""
    ref_start = False
    count = 0
    for para in doc.paragraphs:
        t = para.text.strip()
        # "References" = Chinese for "References"; marks the bibliography section
        # boundary so font superscript conversion only applies to body text.
        if "References" in t and para.style.name.startswith("Heading"):
            ref_start = True
        if ref_start:
            continue
        for run in list(para.runs):
            if _split_run_for_superscript(run):
                count += 1
    return count


# ==================== Main ====================


def main():
    parser = argparse.ArgumentParser(
        description="Convert DOCX reference numbering style",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog=__doc__,
    )
    parser.add_argument("input", help="Input .docx file path")
    parser.add_argument("output", nargs="?", help="Output .docx file path (default: overwrite input)")
    parser.add_argument(
        "--ref-list-style",
        choices=["keep", "bracket", "period"],
        default="keep",
        help="Bibliography list numbering style: keep(no change), bracket([N]), period(N.)",
    )
    parser.add_argument(
        "--style",
        choices=["keep", "unicode", "font"],
        default="keep",
        help="Inline citation style: keep(no change), unicode(Unicode superscript), font(font superscript)",
    )
    parser.add_argument("--range-joiner", default=None, help="Consecutive citation range joiner")
    parser.add_argument("--group-separator", default=None, help="Non-consecutive citation separator")

    args = parser.parse_args()
    src, output = args.input, args.output or args.input

    doc = Document(src)
    summary = []

    # 1) Bibliography list numbering
    if args.ref_list_style == "period":
        c = _convert_ref_list_bracket_to_period(doc)
        summary.append(f"Bibliography: [N]->N. ({c} entries)")
    elif args.ref_list_style == "bracket":
        c = _convert_ref_list_period_to_bracket(doc)
        summary.append(f"Bibliography: N.->[N] ({c} entries)")

    # 2) Inline citations
    if args.style == "unicode":
        if args.range_joiner is None and args.group_separator is None:
            print("Warning: --style unicode but no --range-joiner or --group-separator specified, will use fallback concatenation")
        c = _apply_unicode_superscript(doc, args.range_joiner, args.group_separator)
        summary.append(f"Inline citations: ->Unicode superscript ({c} paragraphs)")
    elif args.style == "font":
        c = _apply_font_superscript(doc)
        summary.append(f"Inline citations: ->Font superscript ({c} runs split)")

    doc.save(output)
    print("Conversion complete.")
    for line in summary:
        print(f"  - {line}")
    print(f"Output file: {output}")


if __name__ == "__main__":
    main()

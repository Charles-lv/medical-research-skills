#!/usr/bin/env python3
"""
generate_docx.py
Convert a Markdown-formatted NSFC grant application to a formatted .docx file.

Usage:
    python generate_docx.py <input.md> <output_dir>

The script reads a Markdown file and produces a Word document with proper
academic formatting for NSFC (National Natural Science Foundation of China) grants.

Dependencies:
    pip install python-docx
"""

import sys
import re
import os
from pathlib import Path
from datetime import datetime


def check_dependencies():
    try:
        from docx import Document
        return True
    except ImportError:
        print("ERROR: python-docx is not installed.")
        print("Install it with: pip install python-docx")
        sys.exit(1)


def setup_document():
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    doc = Document()

    # Page margins (standard academic: top/bottom 2.54cm, left/right 3.17cm)
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.17)
        section.right_margin = Cm(3.17)

    # Default paragraph style: SimSun 12pt, 1.5x line spacing, first-line indent
    from docx.shared import Inches
    from docx.oxml.ns import qn
    normal_style = doc.styles['Normal']
    # 'SimSun' (SimSun) and 'SimHei' (SimHei) are Chinese font family names required
    # for NSFC grant formatting. They must remain as Chinese strings because
    # python-docx uses them as literal font identifiers in the Word XML.
    normal_style.font.name = 'SimSun'
    normal_style.font.size = Pt(12)
    normal_style.element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')

    return doc


def set_font(run, name_cn='SimSun', name_en='Times New Roman', size_pt=12, bold=False):
    """Set font properties on a run.
    name_cn uses Chinese font names ('SimSun'=SimSun, 'SimHei'=SimHei) as these are
    the literal identifiers needed in Word XML for Chinese font rendering."""
    from docx.shared import Pt
    from docx.oxml.ns import qn
    run.font.name = name_en
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.element.rPr.rFonts.set(qn('w:eastAsia'), name_cn)


def set_paragraph_format(para, first_line_indent_chars=0, space_before=0, space_after=0, line_spacing=1.5, alignment=None):
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    pf = para.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)

    # Line spacing
    from docx.shared import Pt as DPt
    pf.line_spacing = Pt(12 * line_spacing)  # 12pt base * multiplier

    if first_line_indent_chars > 0:
        # 2 chars indent = 24pt for 12pt font
        pf.first_line_indent = Pt(12 * first_line_indent_chars)

    if alignment:
        pf.alignment = alignment


def add_heading(doc, text, level):
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn

    para = doc.add_paragraph()
    run = para.add_run(text)

    if level == 1:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(16)
        run.font.bold = True
        run.element.rPr.rFonts.set(qn('w:eastAsia'), 'SimHei')
        para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.paragraph_format.space_before = Pt(12)
        para.paragraph_format.space_after = Pt(6)
    elif level == 2:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(14)
        run.font.bold = True
        run.element.rPr.rFonts.set(qn('w:eastAsia'), 'SimHei')
        para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        para.paragraph_format.space_before = Pt(10)
        para.paragraph_format.space_after = Pt(4)
    elif level == 3:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        run.font.bold = True
        run.element.rPr.rFonts.set(qn('w:eastAsia'), 'SimHei')
        para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        para.paragraph_format.space_before = Pt(8)
        para.paragraph_format.space_after = Pt(2)
    else:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        run.font.bold = True
        run.element.rPr.rFonts.set(qn('w:eastAsia'), 'SimHei')
        para.paragraph_format.space_before = Pt(6)
        para.paragraph_format.space_after = Pt(2)

    return para


def add_body_paragraph(doc, text, is_reference=False):
    from docx.shared import Pt
    from docx.oxml.ns import qn

    para = doc.add_paragraph()

    if is_reference:
        para.paragraph_format.first_line_indent = Pt(-21)
        para.paragraph_format.left_indent = Pt(21)
        para.paragraph_format.space_after = Pt(2)
        para.paragraph_format.line_spacing = Pt(15.75)
        size = Pt(10.5)
    else:
        para.paragraph_format.first_line_indent = Pt(24)
        para.paragraph_format.space_after = Pt(0)
        para.paragraph_format.line_spacing = Pt(21)
        size = Pt(12)

    # Split text on **bold** markers and render each segment
    segments = re.split(r'\*\*(.+?)\*\*', text)
    for idx, seg in enumerate(segments):
        if not seg:
            continue
        run = para.add_run(seg)
        run.font.name = 'Times New Roman'
        run.font.size = size
        run.element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
        if idx % 2 == 1:  # odd segments are inside **...**
            run.font.bold = True
            run.element.rPr.rFonts.set(qn('w:eastAsia'), 'SimHei')

    return para


def parse_and_render(doc, lines):
    """Parse markdown lines and render to docx."""
    i = 0
    in_references = False

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # Skip empty lines (add small spacing)
        if not stripped:
            i += 1
            continue

        # Detect reference section
        # "References" is the Chinese heading for "References" - kept as-is for
        # pattern matching in Chinese-language NSFC grant documents.
        if re.match(r'^References\s*$', stripped):
            in_references = False  # heading itself is normal
            add_heading(doc, stripped, 3)
            in_references = True
            i += 1
            continue

        # Headings
        h4 = re.match(r'^####\s+(.*)', stripped)
        h3 = re.match(r'^###\s+(.*)', stripped)
        h2 = re.match(r'^##\s+(.*)', stripped)
        h1 = re.match(r'^#\s+(.*)', stripped)

        if h1:
            add_heading(doc, h1.group(1), 1)
            in_references = False
        elif h2:
            add_heading(doc, h2.group(1), 2)
            in_references = False
        elif h3:
            add_heading(doc, h3.group(1), 3)
            in_references = False
        elif h4:
            add_heading(doc, h4.group(1), 4)
            in_references = False
        else:
            # Body text
            # Strip leading fullwidth spaces used in template
            text = stripped.lstrip('\u3000')
            if text:
                add_body_paragraph(doc, text, is_reference=in_references)

        i += 1


def main():
    if len(sys.argv) < 3:
        print("Usage: python generate_docx.py <input.md> <output_dir>")
        sys.exit(1)

    check_dependencies()

    input_path = Path(sys.argv[1])
    output_dir = Path(sys.argv[2])

    if not input_path.exists():
        print(f"ERROR: Input file not found: {input_path}")
        sys.exit(1)

    output_dir.mkdir(parents=True, exist_ok=True)

    # Read markdown content
    with open(input_path, 'r', encoding='utf-8') as f:
        content = f.read()

    lines = content.splitlines()

    # Build document
    doc = setup_document()
    parse_and_render(doc, lines)

    # Output filename: same stem as input
    output_path = output_dir / (input_path.stem + '.docx')
    doc.save(str(output_path))
    print(f"SUCCESS: Document saved to {output_path}")


if __name__ == '__main__':
    main()

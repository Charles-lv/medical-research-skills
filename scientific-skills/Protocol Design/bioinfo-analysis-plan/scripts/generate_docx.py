#!/usr/bin/env python3
"""
generate_docx.py
Convert bioinformatics analysis plan Markdown to formatted .docx.

Usage:
    python3 generate_docx.py <input.md> <output_dir>

Produces a Word document with proper formatting for bioinformatics
analysis workflow reports and custom analysis plans.

Dependencies:
    pip install python-docx
"""

import sys
import re
import os
from pathlib import Path


def check_dependencies():
    try:
        from docx import Document
    except ImportError:
        print("ERROR: python-docx is not installed.")
        print("Install it with: pip install python-docx")
        sys.exit(1)


def setup_document():
    from docx import Document
    from docx.shared import Pt, Cm
    from docx.oxml.ns import qn

    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)

    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Times New Roman'
    normal_style.font.size = Pt(11)
    normal_style.element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')

    return doc


def set_run_font(run, bold=False, italic=False, size_pt=11, font_cn='SimSun'):
    from docx.shared import Pt
    from docx.oxml.ns import qn
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.italic = italic
    run.element.rPr.rFonts.set(qn('w:eastAsia'), font_cn)


def add_heading_styled(doc, text, level):
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn

    para = doc.add_paragraph()
    run = para.add_run(text)

    sizes = {1: 18, 2: 15, 3: 13, 4: 11}
    size_pt = sizes.get(level, 11)
    bold = True

    set_run_font(run, bold=True, size_pt=size_pt, font_cn='SimHei')

    if level == 1:
        para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.paragraph_format.space_before = Pt(18)
        para.paragraph_format.space_after = Pt(10)
    elif level == 2:
        para.paragraph_format.space_before = Pt(14)
        para.paragraph_format.space_after = Pt(6)
    elif level == 3:
        para.paragraph_format.space_before = Pt(10)
        para.paragraph_format.space_after = Pt(4)
    else:
        para.paragraph_format.space_before = Pt(6)
        para.paragraph_format.space_after = Pt(3)

    return para


def parse_inline_text(text):
    """Parse inline formatting: **bold**, *italic*, `code`, and markdown links."""
    segments = re.split(r'(\*\*.+?\*\*|\*.+?\*|`[^`]+`|\[.+?\]\(.+?\))', text)
    result = []
    for seg in segments:
        if not seg:
            continue
        bm = re.match(r'\*\*(.+?)\*\*', seg)
        im = re.match(r'\*(.+?)\*', seg)
        cm = re.match(r'`(.+?)`', seg)
        lm = re.match(r'\[(.+?)\]\((.+?)\)', seg)
        if bm:
            result.append(('bold', bm.group(1)))
        elif im:
            result.append(('italic', im.group(1)))
        elif cm:
            result.append(('code', cm.group(1)))
        elif lm:
            result.append(('link', lm.group(1), lm.group(2)))
        else:
            result.append(('text', seg))
    return result


def add_body_paragraph(doc, text, is_reference=False):
    from docx.shared import Pt

    para = doc.add_paragraph()
    segments = parse_inline_text(text)

    if is_reference:
        para.paragraph_format.first_line_indent = Pt(-21)
        para.paragraph_format.left_indent = Pt(21)
        para.paragraph_format.space_after = Pt(2)
        para.paragraph_format.line_spacing = Pt(15.75)
        sz = 10
    else:
        para.paragraph_format.first_line_indent = Pt(22)
        para.paragraph_format.space_after = Pt(2)
        para.paragraph_format.line_spacing = Pt(18)
        sz = 11

    for seg in segments:
        if seg[0] == 'bold':
            run = para.add_run(seg[1])
            set_run_font(run, bold=True, size_pt=sz)
        elif seg[0] == 'italic':
            run = para.add_run(seg[1])
            set_run_font(run, italic=True, size_pt=sz)
        elif seg[0] == 'code':
            run = para.add_run(seg[1])
            set_run_font(run, size_pt=sz, font_cn='Consolas')
        elif seg[0] == 'link':
            run = para.add_run(seg[1])
            set_run_font(run, size_pt=sz)
            run.font.underline = True
            from docx.shared import RGBColor
            run.font.color.rgb = RGBColor(0x05, 0x63, 0xC1)
        else:
            run = para.add_run(seg[1])
            set_run_font(run, size_pt=sz)

    return para


def render_bullet_list(doc, items, ordered=False):
    from docx.shared import Pt
    from docx.oxml.ns import qn

    for item in items:
        para = doc.add_paragraph()
        run = para.add_run(item)
        set_run_font(run, size_pt=11)

        para.paragraph_format.left_indent = Pt(24)
        para.paragraph_format.first_line_indent = Pt(-12)
        para.paragraph_format.space_after = Pt(1)
        para.paragraph_format.line_spacing = Pt(16)

        if ordered:
            num = items.index(item) + 1
            para.insert_paragraph_before()
            run.text = f"  {num}. {item}"
            # Actually let's just do it simpler
            run.text = f"  {items.index(item) + 1}. {item}"
        else:
            run.text = f"  • {item}"


def parse_table_rows(lines):
    """Parse markdown pipe table into header + rows."""
    rows = []
    for line in lines:
        stripped = line.strip()
        if not stripped or stripped.startswith('|---') or stripped.startswith('|:---'):
            continue
        if stripped.startswith('|') and stripped.endswith('|'):
            cells = [c.strip() for c in stripped.split('|')[1:-1]]
            rows.append(cells)
    return rows


def add_table(doc, header_row, data_rows):
    from docx.shared import Pt, Cm, RGBColor
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    if not data_rows:
        return

    num_cols = len(header_row)
    num_rows = len(data_rows) + 1
    table = doc.add_table(rows=num_rows, cols=num_cols)
    table.style = 'Table Grid'
    table.autofit = True

    # Set column widths
    col_width = Cm(14.0 / max(num_cols, 1))

    for i, cell in enumerate(table.columns[0].cells):
        cell.width = Cm(3.0)

    for row_idx, row_data in enumerate([header_row] + data_rows):
        for col_idx, cell_text in enumerate(row_data):
            cell = table.rows[row_idx].cells[col_idx]
            cell.text = ''
            para = cell.paragraphs[0]
            run = para.add_run(cell_text)
            sz = 10
            set_run_font(run, bold=(row_idx == 0), size_pt=sz)

            if row_idx == 0:
                shading = OxmlElement('w:shd')
                shading.set(qn('w:fill'), 'D9E2F3')
                shading.set(qn('w:val'), 'clear')
                cell._tc.get_or_add_tcPr().append(shading)


def render_table(doc, lines, start_idx):
    """Collect all pipe table lines starting from start_idx and render."""
    table_lines = []
    i = start_idx
    while i < len(lines):
        stripped = lines[i].strip()
        if stripped.startswith('|'):
            table_lines.append(lines[i])
            i += 1
        else:
            break

    rows = parse_table_rows(table_lines)
    if len(rows) >= 2:
        add_table(doc, rows[0], rows[1:])
    return i


def render_code_block(doc, lines, start_idx):
    """Render a fenced code block as monospace text."""
    from docx.shared import Pt, Cm
    code_lines = []
    i = start_idx
    if lines[i].strip().startswith('```'):
        i += 1
    while i < len(lines) and not lines[i].strip().startswith('```'):
        code_lines.append(lines[i].rstrip())
        i += 1
    if i < len(lines) and lines[i].strip().startswith('```'):
        i += 1

    if code_lines:
        para = doc.add_paragraph()
        for cl in code_lines:
            run = para.add_run(cl + '\n')
            set_run_font(run, size_pt=9, font_cn='Consolas')
        para.paragraph_format.left_indent = Cm(0.5)
        para.paragraph_format.space_before = Pt(4)
        para.paragraph_format.space_after = Pt(4)
    return i


def parse_and_render(doc, lines):
    from docx.shared import Pt, Cm
    i = 0
    in_references = False
    in_list = False
    list_items = []
    is_ordered_list = False

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if not stripped:
            if in_list and list_items:
                if is_ordered_list:
                    for idx, item in enumerate(list_items):
                        para = doc.add_paragraph()
                        run = para.add_run(f"{idx + 1}. {item}")
                        set_run_font(run, size_pt=11)
                        para.paragraph_format.left_indent = Cm(1.0)
                        para.paragraph_format.space_after = Pt(1)
                else:
                    for item in list_items:
                        para = doc.add_paragraph()
                        run = para.add_run(f"• {item}")
                        set_run_font(run, size_pt=11)
                        para.paragraph_format.left_indent = Cm(1.0)
                        para.paragraph_format.space_after = Pt(1)
                list_items = []
                in_list = False
            i += 1
            continue

        if stripped.startswith('```'):
            i = render_code_block(doc, lines, i)
            continue

        if stripped.startswith('|'):
            if in_list and list_items:
                list_items = []
                in_list = False
            i = render_table(doc, lines, i)
            continue

        if re.match(r'^References\s*$', stripped):
            in_references = False
            add_heading_styled(doc, stripped, 3)
            in_references = True
            i += 1
            continue

        h1 = re.match(r'^#\s+(.*)', stripped)
        h2 = re.match(r'^##\s+(.*)', stripped)
        h3 = re.match(r'^###\s+(.*)', stripped)
        h4 = re.match(r'^####\s+(.*)', stripped)

        if h1:
            if in_list and list_items:
                list_items = []
                in_list = False
            add_heading_styled(doc, h1.group(1), 1)
            in_references = False
        elif h2:
            if in_list and list_items:
                list_items = []
                in_list = False
            add_heading_styled(doc, h2.group(1), 2)
            in_references = False
        elif h3:
            if in_list and list_items:
                list_items = []
                in_list = False
            add_heading_styled(doc, h3.group(1), 3)
            in_references = False
        elif h4:
            if in_list and list_items:
                list_items = []
                in_list = False
            add_heading_styled(doc, h4.group(1), 4)
            in_references = False
        elif stripped.startswith('- ') or stripped.startswith('* '):
            in_list = True
            is_ordered_list = False
            list_items.append(stripped[2:])
        elif re.match(r'^\d+[\.\)]\s', stripped):
            in_list = True
            is_ordered_list = True
            list_items.append(re.sub(r'^\d+[\.\)]\s', '', stripped))
        else:
            if in_list and list_items:
                list_items = []
                in_list = False
            text = stripped.lstrip('\u3000').lstrip('>').strip()
            if text:
                add_body_paragraph(doc, text, is_reference=in_references)

        i += 1

    # Flush remaining list items
    if in_list and list_items:
        if is_ordered_list:
            for idx, item in enumerate(list_items):
                para = doc.add_paragraph()
                run = para.add_run(f"{idx + 1}. {item}")
                set_run_font(run, size_pt=11)
                para.paragraph_format.left_indent = Cm(1.0)
                para.paragraph_format.space_after = Pt(1)
        else:
            for item in list_items:
                para = doc.add_paragraph()
                run = para.add_run(f"• {item}")
                set_run_font(run, size_pt=11)
                para.paragraph_format.left_indent = Cm(1.0)
                para.paragraph_format.space_after = Pt(1)


def main():
    if len(sys.argv) < 3:
        print("Usage: python3 generate_docx.py <input.md> <output_dir>")
        sys.exit(1)

    check_dependencies()

    input_path = Path(sys.argv[1])
    output_dir = Path(sys.argv[2])

    if not input_path.exists():
        print(f"ERROR: Input file not found: {input_path}")
        sys.exit(1)

    output_dir.mkdir(parents=True, exist_ok=True)

    with open(input_path, 'r', encoding='utf-8') as f:
        content = f.read()

    lines = content.splitlines()
    doc = setup_document()
    parse_and_render(doc, lines)

    output_path = output_dir / (input_path.stem + '.docx')
    doc.save(str(output_path))
    print(f"SUCCESS: Document saved to {output_path}")


if __name__ == '__main__':
    main()
